from pathlib import Path
from langchain_core.documents import Document


class DocumentLoader:
    SUFFIX_MAP = {
        ".md": "markdown", ".txt": "text",
        ".pdf": "pdf", ".docx": "docx",
    }

    async def load(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix not in self.SUFFIX_MAP:
            raise ValueError(f"Unsupported format: {suffix}")

        if suffix in (".md", ".txt"):
            return await self._load_text(path)
        elif suffix == ".pdf":
            return await self._load_pdf(path)
        elif suffix == ".docx":
            return await self._load_docx(path)
        return []

    async def load_directory(self, dir_path: str) -> list[Document]:
        docs = []
        for file_path in Path(dir_path).glob("*"):
            if file_path.suffix.lower() in self.SUFFIX_MAP:
                docs.extend(await self.load(str(file_path)))
        return docs

    async def _load_text(self, path: Path) -> list[Document]:
        with open(path, encoding="utf-8") as f:
            text = f.read()
        meta = {"doc_name": path.name, "doc_type": path.suffix.lstrip(".")}
        return [Document(page_content=text, metadata=meta)]

    async def _load_pdf(self, path: Path) -> list[Document]:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"doc_name": path.name, "page": i + 1, "doc_type": "pdf"},
                ))
        return docs

    async def _load_docx(self, path: Path) -> list[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=text, metadata={"doc_name": path.name, "doc_type": "docx"})]
